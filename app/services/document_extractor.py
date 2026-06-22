import io
import re
from pathlib import Path

import fitz
from docx import Document
from fastapi import HTTPException, UploadFile, status

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


async def read_upload(file: UploadFile, max_bytes: int) -> bytes:
    data = await file.read()
    if len(data) > max_bytes:
        raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="Resume exceeds 10 MB limit")
    return data


def validate_upload(file: UploadFile) -> str:
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Only PDF, DOCX, and TXT resumes are supported")
    return suffix


def extract_text(data: bytes, suffix: str) -> str:
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix == ".txt":
        return data.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file extension: {suffix}")


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"Page\s+\d+\s+(of\s+\d+)?", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"[\r\t]+", " ", text)
    text = re.sub(r"[^\S\n]+", " ", text)
    text = text.replace(" \n", "\n")
    text = text.replace("\n ", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[^\w\s@.+:/,#()&%-]", " ", text)
    return re.sub(r" {2,}", " ", text).strip()


def _extract_pdf(data: bytes) -> str:
    chunks: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            chunks.append(page.get_text("text"))
    return "\n".join(chunks)


def _extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.append(" ".join(cell.text for cell in row.cells))
    return "\n".join([*paragraphs, *table_text])
